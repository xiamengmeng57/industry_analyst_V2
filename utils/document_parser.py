"""
文档解析工具 - 支持多种格式的文档解析
"""
import os
import json
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime


class DocumentParser:
    """文档解析器 - 支持 PDF、DOCX、TXT、JSON 等格式"""

    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt', '.md', '.json']

    @staticmethod
    def parse_document(file_path: str) -> Dict[str, Any]:
        """
        解析文档并提取内容

        Args:
            file_path: 文档文件路径

        Returns:
            包含文档内容和元数据的字典
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文档不存在: {file_path}")

        file_path_obj = Path(file_path)
        file_ext = file_path_obj.suffix.lower()

        if file_ext not in DocumentParser.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文档格式: {file_ext}. 支持的格式: {DocumentParser.SUPPORTED_FORMATS}")

        # 根据文件类型选择解析方法
        if file_ext == '.pdf':
            content = DocumentParser._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            content = DocumentParser._parse_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            content = DocumentParser._parse_text(file_path)
        elif file_ext == '.json':
            content = DocumentParser._parse_json(file_path)
        else:
            raise ValueError(f"暂未实现的格式: {file_ext}")

        # 提取元数据
        metadata = {
            "filename": file_path_obj.name,
            "file_path": str(file_path_obj.absolute()),
            "file_size": os.path.getsize(file_path),
            "file_type": file_ext,
            "parsed_at": datetime.now().isoformat()
        }

        return {
            "content": content,
            "metadata": metadata
        }

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """解析 PDF 文件"""
        try:
            import PyPDF2

            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)

            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("需要安装 PyPDF2: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"PDF 解析失败: {e}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """解析 DOCX 文件"""
        try:
            import docx

            doc = docx.Document(file_path)
            text_parts = []

            # 提取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"DOCX 解析失败: {e}")

    @staticmethod
    def _parse_text(file_path: str) -> str:
        """解析纯文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read()
            except Exception as e:
                raise Exception(f"文本文件解析失败: {e}")
        except Exception as e:
            raise Exception(f"文本文件解析失败: {e}")

    @staticmethod
    def _parse_json(file_path: str) -> Any:
        """
        解析 JSON 文件

        返回:
        - 如果 JSON 是对象或数组，返回原始结构（保留为 dict/list）
        - 如果是其他类型，转换为字符串
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                # 如果是 dict 或 list，保持原样返回
                # 否则转为字符串
                if isinstance(data, (dict, list)):
                    return data
                else:
                    return str(data)
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    data = json.load(file)
                    return data if isinstance(data, (dict, list)) else str(data)
            except Exception as e:
                raise Exception(f"JSON 文件解析失败: {e}")
        except json.JSONDecodeError as e:
            raise Exception(f"JSON 格式错误: {e}")
        except Exception as e:
            raise Exception(f"JSON 文件解析失败: {e}")

    @staticmethod
    def convert_to_findings(
        file_path: str,
        title: Optional[str] = None,
        custom_metadata: Optional[Dict] = None,
        extract_time: bool = False,
        anthropic_client=None
    ) -> Dict[str, Any]:
        """
        将上传的文档转换为 findings 格式

        Args:
            file_path: 文档路径
            title: 自定义标题（可选，默认使用文件名）
            custom_metadata: 自定义元数据（可选）
            extract_time: 是否提取时间事实（默认 False）
            anthropic_client: Anthropic client（提取时间时必需）

        Returns:
            符合 findings 格式的字典
        """
        # 解析文档
        parsed = DocumentParser.parse_document(file_path)
        content = parsed["content"]
        doc_metadata = parsed["metadata"]

        # 使用文件名作为标题（如果未提供）
        if title is None:
            title = doc_metadata["filename"]

        # JSON 格式特殊处理
        if doc_metadata["file_type"] == '.json':
            # 如果内容是 dict 或 list，转换为格式化的字符串
            if isinstance(content, (dict, list)):
                content = json.dumps(content, ensure_ascii=False, indent=2)

        # 提取时间事实（如果启用）
        publication_date = None
        time_facts = []
        if extract_time and anthropic_client:
            print(f"  📅 提取时间事实: {doc_metadata['filename']}")
            time_info = DocumentParser.extract_time_facts(
                content,
                client=anthropic_client
            )
            publication_date = time_info.get("publication_date")
            time_facts = time_info.get("time_facts", [])

        # 构建 findings 格式
        finding = {
            "topic": title,
            "data": content,
            "source": f"user_upload://{doc_metadata['file_path']}",
            "date": publication_date or datetime.now().strftime("%Y-%m-%d"),  # 优先使用提取的发布时间
            "publication_date": publication_date or datetime.now().strftime("%Y-%m-%d"),
            "time_facts": time_facts,  # 提取的时间-事实对
            "metadata": {
                "source_type": "user_upload",
                "is_user_provided": True,
                "original_filename": doc_metadata["filename"],
                "file_type": doc_metadata["file_type"],
                "file_size": doc_metadata["file_size"],
                "parsed_at": doc_metadata["parsed_at"],
                "time_extracted": extract_time  # 标记是否进行了时间提取
            }
        }

        # 合并自定义元数据
        if custom_metadata:
            finding["metadata"].update(custom_metadata)

        return finding

    @staticmethod
    def convert_to_verified(
        file_path: str,
        title: Optional[str] = None,
        confidence: float = 0.95,
        custom_metadata: Optional[Dict] = None,
        extract_time: bool = False,
        anthropic_client=None
    ) -> Dict[str, Any]:
        """
        将上传的文档直接转换为 verified 格式

        Args:
            file_path: 文档路径
            title: 自定义标题（可选）
            confidence: 置信度（默认 0.95，表示用户提供的可信文档）
            custom_metadata: 自定义元数据（可选）
            extract_time: 是否提取时间事实（默认 False）
            anthropic_client: Anthropic client（提取时间时必需）

        Returns:
            符合 verified 格式的字典
        """
        # 先转换为 findings 格式（包含时间提取）
        finding = DocumentParser.convert_to_findings(
            file_path,
            title,
            custom_metadata,
            extract_time=extract_time,
            anthropic_client=anthropic_client
        )

        # 转换为 verified 格式
        verified = {
            "fact": finding["data"],
            "time": finding["date"],
            "publication_date": finding["publication_date"],
            "confidence": confidence,
            "source": finding["source"],
            "supporting_sources": [finding["source"]],
            "cross_verified": False,  # 单一来源（用户上传）
            "time_verified": True,  # 假设用户文档的时间是准确的
            "timeliness": "用户提供",
            "time_conflicts": [],
            "metadata": finding["metadata"]
        }

        return verified

    @staticmethod
    def batch_convert_to_findings(
        file_paths: List[str],
        titles: Optional[List[str]] = None,
        extract_time: bool = False,
        anthropic_client=None
    ) -> List[Dict[str, Any]]:
        """
        批量转换多个文档为 findings 格式

        Args:
            file_paths: 文档路径列表
            titles: 标题列表（可选，需与 file_paths 长度一致）
            extract_time: 是否提取时间事实（默认 False）
            anthropic_client: Anthropic client（提取时间时必需）

        Returns:
            findings 列表
        """
        if titles and len(titles) != len(file_paths):
            raise ValueError("titles 长度必须与 file_paths 一致")

        findings = []
        for i, file_path in enumerate(file_paths):
            try:
                title = titles[i] if titles else None
                finding = DocumentParser.convert_to_findings(
                    file_path,
                    title,
                    extract_time=extract_time,
                    anthropic_client=anthropic_client
                )
                findings.append(finding)
            except Exception as e:
                print(f"警告: 解析文档 {file_path} 失败: {e}")
                continue

        return findings

    @staticmethod
    def batch_convert_to_verified(
        file_paths: List[str],
        titles: Optional[List[str]] = None,
        confidence: float = 0.95,
        extract_time: bool = False,
        anthropic_client=None
    ) -> List[Dict[str, Any]]:
        """
        批量转换多个文档为 verified 格式

        Args:
            file_paths: 文档路径列表
            titles: 标题列表（可选）
            confidence: 置信度
            extract_time: 是否提取时间事实（默认 False）
            anthropic_client: Anthropic client（提取时间时必需）

        Returns:
            verified 列表
        """
        if titles and len(titles) != len(file_paths):
            raise ValueError("titles 长度必须与 file_paths 一致")

        verified_list = []
        for i, file_path in enumerate(file_paths):
            try:
                title = titles[i] if titles else None
                verified = DocumentParser.convert_to_verified(
                    file_path,
                    title,
                    confidence,
                    extract_time=extract_time,
                    anthropic_client=anthropic_client
                )
                verified_list.append(verified)
            except Exception as e:
                print(f"警告: 解析文档 {file_path} 失败: {e}")
                continue

        return verified_list

    @staticmethod
    def extract_time_facts(
        content: str,
        client=None,
        model: str = "claude-haiku-4-5-20251001"
    ) -> Dict[str, Any]:
        """
        使用 LLM 提取文本中的时间和关联事实

        （复用 Researcher Agent 的时间提取逻辑）

        Args:
            content: 文本内容
            client: Anthropic client（必需）
            model: 使用的模型

        Returns:
            包含 publication_date, time_facts 的字典
            time_facts 格式: [{"fact": "事实描述", "time": "时间", "time_type": "mentioned"}]
        """
        if client is None:
            # 如果没有提供 client，返回空结果
            print("  ⚠️  未提供 Anthropic client，跳过时间提取")
            return {"publication_date": None, "time_facts": []}

        if not content or len(content.strip()) < 20:
            return {"publication_date": None, "time_facts": []}

        try:
            prompt = f"""请从以下文本中提取**所有**时间信息和关联事实：

1. **发布时间**：这篇内容的发布/创建时间（如果文中有提到）
2. **时间-事实对**：文中提到的**每个**具体事实及其对应时间

要求：
- 提取所有事件和时间对，一个文本中可能有多个事件和多个时间
- 时间格式：保持原文格式（如：2024年3月、2024-03-15、2024Q1、近期）
- 事实描述：简洁准确，一句话概括每个独立事件
- 时间类型：
  - "publication"：内容发布时间
  - "mentioned"：文中明确提到的时间
- 只提取明确的信息，不要推测
- 如果一个事件关联多个时间，创建多个条目
- 如果一个时间关联多个事件，创建多个条目

返回JSON格式：
{{
  "publication_date": "发布时间或null",
  "time_facts": [
    {{"fact": "事件1描述", "time": "2024年3月", "time_type": "mentioned"}},
    {{"fact": "事件2描述", "time": "2024年Q1", "time_type": "mentioned"}},
    {{"fact": "事件3描述", "time": "2024-03-15", "time_type": "mentioned"}}
  ]
}}

文本内容（前2000字）：
{content[:2000]}

请返回JSON格式的结果，确保提取所有事件-时间对。"""

            response = client.messages.create(
                model=model,
                max_tokens=1024,  # 增加 token 以支持更长文档
                temperature=0,
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )

            # 提取文本内容（跳过 ThinkingBlock）
            text_parts = []
            for block in response.content:
                if hasattr(block, 'text'):
                    text_parts.append(block.text)
            result_text = '\n'.join(text_parts)

            # 提取 JSON
            if "{" in result_text and "}" in result_text:
                start = result_text.find("{")
                end = result_text.rfind("}") + 1
                json_text = result_text[start:end]
                result = json.loads(json_text)
            else:
                result = json.loads(result_text)

            return {
                "publication_date": result.get("publication_date"),
                "time_facts": result.get("time_facts", [])
            }

        except Exception as e:
            print(f"  ⚠️ 提取时间事实失败: {str(e)}")
            return {"publication_date": None, "time_facts": []}
